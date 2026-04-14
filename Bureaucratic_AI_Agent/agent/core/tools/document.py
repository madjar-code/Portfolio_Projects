import asyncio
import io
import logging
import re
from core.tools import Tool, ToolRegistry
from core.tools._fetcher import DocumentFetcher
from core.tools._vision import vision_extract
from models import TaskMessage

logger = logging.getLogger(__name__)

_OCR_PROMPT = "Extract all text from this document image. Return plain text only, preserving layout."
_REGION_PROMPT = "Extract all text from this image region. Return plain text only."
_DOCX_CHUNK_SIZE = 30  # paragraphs per "page"


# ---------------------------------------------------------------------------
# Internal readers
# ---------------------------------------------------------------------------

async def _read_pdf_page(content: bytes, page_number: int) -> str:
    import pdfplumber
    import fitz  # pymupdf

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        total = len(pdf.pages)
        if page_number < 1 or page_number > total:
            return f"Page {page_number} does not exist (document has {total} pages)."
        text = pdf.pages[page_number - 1].extract_text() or ""

    if text.strip():
        return text

    # Scanned PDF — render page as image and send to Vision
    logger.debug("PDF page %d has no text layer, falling back to Vision OCR", page_number)
    doc = fitz.open(stream=content, filetype="pdf")
    pix = doc[page_number - 1].get_pixmap(dpi=150)
    png_bytes = pix.tobytes("png")
    return await vision_extract(png_bytes, "image/png", _OCR_PROMPT)


async def _read_docx_page(content: bytes, page_number: int) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    start = (page_number - 1) * _DOCX_CHUNK_SIZE
    chunk = paras[start : start + _DOCX_CHUNK_SIZE]
    if not chunk:
        return f"Page {page_number} does not exist."
    return "\n".join(chunk)


async def _count_pages(content: bytes, file_format: str) -> int:
    if file_format == "PDF":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return len(pdf.pages)
    elif file_format == "DOCX":
        from docx import Document
        doc = Document(io.BytesIO(content))
        paras = [p for p in doc.paragraphs if p.text.strip()]
        return max(1, (len(paras) + _DOCX_CHUNK_SIZE - 1) // _DOCX_CHUNK_SIZE)
    else:  # JPG / PNG
        return 1


# ---------------------------------------------------------------------------
# Tool functions (called by the registry closures)
# ---------------------------------------------------------------------------

async def _read_document_page(
    url: str, file_format: str, fetcher: DocumentFetcher, page_number: int
) -> str:
    content = await fetcher.fetch(url)

    if file_format == "PDF":
        return await _read_pdf_page(content, page_number)
    elif file_format == "DOCX":
        return await _read_docx_page(content, page_number)
    elif file_format in ("JPG", "PNG"):
        mime = "image/jpeg" if file_format == "JPG" else "image/png"
        return await vision_extract(content, mime, _OCR_PROMPT)
    return f"Unsupported format: {file_format}"


async def _ocr_document_region(
    url: str, file_format: str, fetcher: DocumentFetcher,
    page: int, x: int, y: int, width: int, height: int,
) -> str:
    from PIL import Image

    content = await fetcher.fetch(url)

    if file_format == "PDF":
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        pix = doc[page - 1].get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
    elif file_format in ("JPG", "PNG"):
        img = Image.open(io.BytesIO(content))
    else:
        return f"ocr_document_region is not supported for format: {file_format}"

    region = img.crop((x, y, x + width, y + height))
    buf = io.BytesIO()
    region.save(buf, format="PNG")
    return await vision_extract(buf.getvalue(), "image/png", _REGION_PROMPT)


async def _extract_field_from_document(
    url: str, file_format: str, fetcher: DocumentFetcher, field_name: str
) -> str:
    content = await fetcher.fetch(url)
    total = await _count_pages(content, file_format)

    pages = await asyncio.gather(*[
        _read_document_page(url, file_format, fetcher, n)
        for n in range(1, total + 1)
    ])
    full_text = "\n".join(pages)

    normalized = field_name.replace("_", " ")
    match = re.search(rf"(?i){re.escape(normalized)}\s*[:\-]\s*(.+)", full_text)
    if match:
        return match.group(1).strip()
    return f"Field '{field_name}' not found in document."


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------

def build_document_registry(task: TaskMessage) -> ToolRegistry:
    registry = ToolRegistry()

    if task.document is None:
        return registry  # no document tools — agent reads form_data from context

    url = task.document.file_url
    fmt = task.document.file_format
    fetcher = DocumentFetcher()

    registry.register(Tool(
        name="read_document_page",
        description="Read a full page of the submitted document and return its text content.",
        parameters={
            "type": "object",
            "properties": {"page_number": {"type": "integer", "description": "1-based page number"}},
            "required": ["page_number"],
        },
        fn=lambda page_number: _read_document_page(url, fmt, fetcher, page_number),
        untrusted=True,
    ))

    registry.register(Tool(
        name="ocr_document_region",
        description="Run OCR on a specific rectangular region of a document page.",
        parameters={
            "type": "object",
            "properties": {
                "page": {"type": "integer"},
                "x": {"type": "integer"},
                "y": {"type": "integer"},
                "width": {"type": "integer"},
                "height": {"type": "integer"},
            },
            "required": ["page", "x", "y", "width", "height"],
        },
        fn=lambda page, x, y, width, height: _ocr_document_region(url, fmt, fetcher, page, x, y, width, height),
        untrusted=True,
    ))

    registry.register(Tool(
        name="extract_field_from_document",
        description="Extract a named field (e.g. first_name, date_of_birth) from the document.",
        parameters={
            "type": "object",
            "properties": {"field_name": {"type": "string"}},
            "required": ["field_name"],
        },
        fn=lambda field_name: _extract_field_from_document(url, fmt, fetcher, field_name),
        untrusted=True,
    ))

    return registry
