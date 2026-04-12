"""
Generate DOCX fixture files for lease_agreement_md eval cases.
Run from the agent/ directory:
    uv run python -m evals.fixtures.lease_agreement_md.generate
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

_OUT = Path(__file__).parent

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True


def _table_row(table, label: str, value: str) -> None:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value


def _save(doc: Document, filename: str) -> None:
    path = _OUT / filename
    doc.save(path)
    size_kb = path.stat().st_size // 1024
    print(f"  OK {filename}  ({size_kb} KB)")


# ---------------------------------------------------------------------------
# Shared content blocks
# ---------------------------------------------------------------------------

_PARTIES_TEXT = (
    "LANDLORD: Ion Rusu, born 15/03/1968, ID series A, No. 1234567, "
    "domiciled at str. Independentei 5, ap. 2, Chisinau, Republic of Moldova.\n\n"
    "TENANT: Maria Cojocaru, born 22/07/1990, ID series B, No. 7654321, "
    "domiciled at str. Florilor 12, ap. 8, Balti, Republic of Moldova."
)

_PROPERTY_TEXT = (
    "The landlord agrees to lease the residential apartment located at "
    "str. Puskin 34, ap. 6, Chisinau, Republic of Moldova, "
    "with a total usable area of 48 m², consisting of 2 rooms, a kitchen, "
    "a bathroom, and a hallway."
)

_DURATION_TEXT = (
    "This agreement shall enter into force on 01/05/2026 and shall remain in effect "
    "until 30/04/2027 (12 months). Upon expiry, the agreement may be renewed by "
    "mutual written consent of both parties."
)

_RENT_TEXT = (
    "The monthly rent is 4,500 MDL (four thousand five hundred Moldovan Lei), "
    "payable no later than the 5th day of each calendar month. Payment shall be "
    "made by bank transfer to the landlord's account or in cash against receipt."
)

_DEPOSIT_TEXT = (
    "Upon signing this agreement, the tenant shall pay a security deposit of "
    "4,500 MDL (equivalent to one month's rent). The deposit shall be returned "
    "within 14 days after the end of the tenancy, minus any deductions for damages "
    "beyond normal wear and tear."
)

_UTILITIES_TEXT = (
    "The tenant is responsible for all utility costs: electricity, natural gas, "
    "cold and hot water, internet, and telephone. Monthly utility invoices shall be "
    "paid directly by the tenant to the respective service providers. "
    "Building maintenance fees (fond de reparatie) are the responsibility of the landlord."
)

_TERMINATION_TEXT = (
    "Either party may terminate this agreement before the expiry date by providing "
    "30 (thirty) days written notice. In the event of material breach of contract "
    "terms by either party, the non-breaching party may terminate with 7 days notice. "
    "Grounds for immediate termination include non-payment of rent for 2 consecutive "
    "months, subletting without consent, or causing significant damage to the property."
)

_GENERAL_CLAUSES = [
    (
        "Condition of Property",
        "The landlord warrants that the property is in good habitable condition at the "
        "start of the tenancy. An inventory of furnishings and appliances is attached as "
        "Annex 1. The tenant agrees to maintain the property in the same condition, "
        "subject to normal wear and tear."
    ),
    (
        "Access",
        "The landlord may access the property only with prior notice of at least 24 hours, "
        "except in cases of emergency. Inspections shall be conducted at mutually agreed "
        "times, no more than once per month."
    ),
    (
        "Subletting",
        "The tenant may not sublet the property or any part thereof without the prior "
        "written consent of the landlord. Violation of this clause shall constitute grounds "
        "for immediate termination of the agreement."
    ),
    (
        "Modifications",
        "The tenant may not make structural modifications, install fixtures, or carry out "
        "renovations without the landlord's prior written consent. Minor decorative changes "
        "(e.g., picture hanging) are permitted."
    ),
    (
        "Disputes",
        "Any disputes arising from or related to this agreement shall be resolved through "
        "negotiation in good faith. If resolution cannot be reached, disputes shall be "
        "submitted to the competent courts of the Republic of Moldova in accordance with "
        "applicable law."
    ),
    (
        "Governing Law",
        "This agreement is governed by the Civil Code of the Republic of Moldova, Law No. 1107 "
        "of 06/06/2002, and other applicable legislation regarding residential tenancy."
    ),
]

_SIGNATURES_TEXT = (
    "This agreement has been concluded in 2 (two) original copies, one for each party, "
    "and shall enter into force upon signature by both parties."
)


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def _base_lease(doc: Document, include: set[str]) -> Document:
    """Build a lease document, omitting sections not in `include`."""

    _heading(doc, "CONTRACT DE LOCATIUNE / RESIDENTIAL LEASE AGREEMENT", level=1)
    _para(doc, "No. 2026-04/001  |  Date: 10 April 2026")
    doc.add_paragraph()

    if "parties" in include:
        _heading(doc, "1. PARTIES", level=2)
        _para(doc, _PARTIES_TEXT)
        doc.add_paragraph()

    if "property" in include:
        _heading(doc, "2. LEASED PROPERTY", level=2)
        _para(doc, _PROPERTY_TEXT)
        doc.add_paragraph()

    if "duration" in include:
        _heading(doc, "3. LEASE DURATION", level=2)
        _para(doc, _DURATION_TEXT)
        doc.add_paragraph()

    if "rent" in include:
        _heading(doc, "4. RENT", level=2)
        _para(doc, _RENT_TEXT)
        doc.add_paragraph()

    if "deposit" in include:
        _heading(doc, "5. SECURITY DEPOSIT", level=2)
        _para(doc, _DEPOSIT_TEXT)
        doc.add_paragraph()

    if "utilities" in include:
        _heading(doc, "6. UTILITIES AND EXPENSES", level=2)
        _para(doc, _UTILITIES_TEXT)
        doc.add_paragraph()

    if "termination" in include:
        _heading(doc, "7. TERMINATION", level=2)
        _para(doc, _TERMINATION_TEXT)
        doc.add_paragraph()

    # General clauses always included for page length
    _heading(doc, "8. GENERAL PROVISIONS", level=2)
    for title, body in _GENERAL_CLAUSES:
        _heading(doc, title, level=3)
        _para(doc, body)
        doc.add_paragraph()

    if "signatures" in include:
        _heading(doc, "9. SIGNATURES", level=2)
        _para(doc, _SIGNATURES_TEXT)
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _table_row(table, "LANDLORD", "TENANT")
        _table_row(table, "Ion Rusu", "Maria Cojocaru")
        _table_row(table, "Signature: _______________", "Signature: _______________")
        _table_row(table, "Date: _______________", "Date: _______________")

    return doc


_ALL_SECTIONS = {"parties", "property", "duration", "rent", "deposit", "utilities", "termination", "signatures"}


def make_valid(filename: str) -> None:
    doc = Document()
    _base_lease(doc, include=_ALL_SECTIONS)
    _save(doc, filename)


def make_no_deposit(filename: str) -> None:
    doc = Document()
    _base_lease(doc, include=_ALL_SECTIONS - {"deposit"})
    _save(doc, filename)


def make_no_signatures(filename: str) -> None:
    doc = Document()
    _base_lease(doc, include=_ALL_SECTIONS - {"signatures"})
    _save(doc, filename)


def make_no_rent(filename: str) -> None:
    doc = Document()
    sections = _ALL_SECTIONS - {"rent"}
    _base_lease(doc, include=sections)
    # Add a rent section that exists but has no amount
    _heading(doc, "4. RENT", level=2)
    _para(doc, "The monthly rent shall be agreed upon separately by both parties.")
    _save(doc, filename)


def make_wrong_document(filename: str) -> None:
    """A company memo — clearly not a lease agreement."""
    doc = Document()
    _heading(doc, "INTERNAL MEMORANDUM", level=1)
    _para(doc, "To: All Department Heads")
    _para(doc, "From: CEO Office")
    _para(doc, "Date: 10 April 2026")
    _para(doc, "Subject: Q1 2026 Performance Review")
    doc.add_paragraph()
    _para(doc, (
        "This memorandum summarizes the outcomes of the Q1 2026 performance review "
        "conducted across all departments. We are pleased to report that overall "
        "targets have been met or exceeded in 7 out of 9 business units."
    ))
    for i in range(1, 6):
        _heading(doc, f"Department {i}: Summary", level=2)
        _para(doc, (
            f"Department {i} achieved {85 + i}% of its quarterly targets. "
            "Key accomplishments include process improvements, cost reductions, "
            "and successful delivery of two major projects. Areas for improvement "
            "identified include response time to client inquiries and inter-team "
            "communication. Action plans have been submitted and will be reviewed "
            "at the next quarterly meeting."
        ))
        doc.add_paragraph()
    _heading(doc, "Next Steps", level=2)
    _para(doc, (
        "Each department head is requested to submit a detailed Q2 action plan by "
        "30 April 2026. The plans will be reviewed by the executive committee and "
        "feedback will be provided within 10 business days."
    ))
    _save(doc, filename)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating lease_agreement_md fixtures...")
    make_valid("lease_valid.docx")
    make_no_deposit("lease_no_deposit.docx")
    make_no_signatures("lease_no_signatures.docx")
    make_no_rent("lease_no_rent.docx")
    make_wrong_document("random_doc.docx")
    print("Done.")
